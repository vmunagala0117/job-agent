"""Utilities to render application package artifacts and assemble ZIP archives.

Generates TXT/MD content and optional DOCX if python-docx is installed.
Produces an in-memory ZIP (bytes) suitable for streaming via FastAPI.
"""

from __future__ import annotations

import io
import zipfile
import datetime
import re
from typing import Iterable, List

def _sanitize_filename(name: str) -> str:
    return re.sub(r"[^A-Za-z0-9 _\-\.]+", "_", name)[:200]


def render_cover_letter_txt(pkg) -> str:
    return pkg.cover_letter or ""


def render_intro_email_txt(pkg) -> str:
    return pkg.intro_email or ""


def render_resume_suggestions_txt(pkg) -> str:
    if not pkg.resume_suggestions:
        return ""
    return "\n".join([f"- {s}" for s in pkg.resume_suggestions])


def render_job_description_txt(pkg) -> str:
    try:
        return pkg.job.description or ""
    except Exception:
        return ""


def render_cover_letter_md(pkg) -> str:
    # Simple markdown wrapper
    title = pkg.job.title if getattr(pkg.job, "title", None) else "Cover Letter"
    header = f"# Cover Letter — {title}\n\n"
    return header + (pkg.cover_letter or "")


def _render_docx_bytes_from_text(text: str, title: str = "Document") -> bytes | None:
    try:
        from docx import Document
    except Exception:
        return None

    doc = Document()
    doc.add_heading(title, level=1)
    for line in text.splitlines():
        if line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


def packages_to_zip_bytes(packages: Iterable, formats: List[str] | None = None) -> bytes:
    """Return ZIP bytes containing artifacts for given packages.

    formats: list containing any of 'txt', 'md', 'docx'. Default ['txt','md']
    """
    if formats is None:
        formats = ["txt", "md"]

    buf = io.BytesIO()
    errors = []
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for pkg in packages:
            try:
                # Create folder name per package
                job_title = getattr(pkg.job, "title", "job") or "job"
                folder = _sanitize_filename(f"{job_title}_{pkg.id[:8]}")

                # Cover letter
                cover_txt = render_cover_letter_txt(pkg)
                if "txt" in formats:
                    zf.writestr(f"{folder}/cover_letter.txt", cover_txt)
                if "md" in formats:
                    zf.writestr(f"{folder}/cover_letter.md", render_cover_letter_md(pkg))
                if "docx" in formats:
                    docb = _render_docx_bytes_from_text(cover_txt, title="Cover Letter")
                    if docb:
                        zf.writestr(f"{folder}/cover_letter.docx", docb)

                # Intro email
                intro_txt = render_intro_email_txt(pkg)
                if "txt" in formats:
                    zf.writestr(f"{folder}/intro_email.txt", intro_txt)
                if "docx" in formats:
                    docb = _render_docx_bytes_from_text(intro_txt, title="Intro Email")
                    if docb:
                        zf.writestr(f"{folder}/intro_email.docx", docb)

                # Resume suggestions (plain text and improved Markdown)
                rs_list = pkg.resume_suggestions or []
                if rs_list:
                    # Plain text (legacy)
                    zf.writestr(f"{folder}/resume_suggestions.txt", "\n\n".join(rs_list))

                    # Markdown rendering: parse LLM blocks into readable sections
                    md_lines = ["# Resume Suggestions", ""]
                    for block in rs_list:
                        # Normalize line endings and strip
                        blk = block.strip()
                        if not blk:
                            continue

                        # If this block looks like a Gap Analysis, render specially
                        if blk.lower().startswith("gap analysis"):
                            md_lines.append("## Gap Analysis")
                            for l in blk.splitlines()[1:]:
                                l = l.strip()
                                if not l:
                                    continue
                                # Ensure bullet formatting
                                if not l.startswith("-"):
                                    l = "- " + l
                                md_lines.append(l)
                            md_lines.append("")
                            continue

                        # Generic suggestion block parsing: find labeled fields
                        lines = [l.rstrip() for l in blk.splitlines() if l.strip()]
                        section = None
                        original = []
                        suggested = []
                        why = []

                        cur = None
                        for ln in lines:
                            low = ln.lower()
                            if ln.startswith("[") and "]" in ln:
                                # [SECTION] on first line
                                section = ln.strip().lstrip("[").split("]")[0].strip()
                                # capture remainder of line after ] if any
                                rem = ln.split("]", 1)[1].strip()
                                if rem:
                                    # if remainder begins with 'Original:' move to that state
                                    if rem.lower().startswith("original:"):
                                        cur = 'original'
                                        original.append(rem.split(':', 1)[1].strip())
                                    else:
                                        # treat as part of original
                                        cur = 'original'
                                        original.append(rem)
                                continue
                            if low.startswith("original:"):
                                cur = 'original'
                                original.append(ln.split(":", 1)[1].strip())
                                continue
                            if low.startswith("suggested change:") or low.startswith("suggested:"):
                                cur = 'suggested'
                                suggested.append(ln.split(":", 1)[1].strip())
                                continue
                            if "- why:" in low or low.startswith("why:") or low.startswith("- why"):
                                cur = 'why'
                                # remove leading markers
                                cleaned = re.sub(r"^[-\s]*why:??\s*", "", ln, flags=re.I)
                                why.append(cleaned.strip())
                                continue

                            # Append to current section if set
                            if cur == 'original':
                                original.append(ln)
                            elif cur == 'suggested':
                                suggested.append(ln)
                            elif cur == 'why':
                                why.append(ln)
                            else:
                                # If we haven't seen a label, attempt to infer: first line -> section
                                if section is None:
                                    m = re.match(r"^\[?([^\]]+)\]?", ln)
                                    if m:
                                        section = m.group(1).strip()
                                        continue
                                # default to treating as original
                                original.append(ln)

                        # Fallback section name
                        if not section:
                            section = "Suggestion"

                        md_lines.append(f"### {section}")
                        if original:
                            md_lines.append("**Original:**")
                            for o in original:
                                md_lines.append(f"> {o}")
                        if suggested:
                            md_lines.append("")
                            md_lines.append("**Suggested Change:**")
                            md_lines.append("```text")
                            for s in suggested:
                                md_lines.append(s)
                            md_lines.append("```")
                        if why:
                            md_lines.append("")
                            md_lines.append("**Why:**")
                            for w in why:
                                md_lines.append(f"- {w}")

                        md_lines.append("")

                    zf.writestr(f"{folder}/resume_suggestions.md", "\n".join(md_lines))

                # Job description
                jd = render_job_description_txt(pkg)
                if jd:
                    zf.writestr(f"{folder}/job_description.txt", jd)

                # Metadata file: include additional job attributes when present
                job = getattr(pkg, "job", None)
                profile_name = getattr(getattr(pkg, "profile", None), "name", "")
                job_title = getattr(job, "title", "") if job is not None else ""
                company = (getattr(job, "company", "") or getattr(job, "company_name", "")) if job is not None else ""
                location = getattr(job, "location", "") if job is not None else ""

                # Posting date: support datetime-like or string values
                posting_date_val = None
                if job is not None:
                    posting_date_val = getattr(job, "posted_at", None) or getattr(job, "posting_date", None) or getattr(job, "posted_date", None)
                posting_date = ""
                if posting_date_val is not None:
                    try:
                        posting_date = posting_date_val.isoformat()
                    except Exception:
                        posting_date = str(posting_date_val)

                # Salary and job URL (try a few common attribute names)
                salary = ""
                if job is not None:
                    salary = getattr(job, "salary_range", "") or getattr(job, "salary", "") or getattr(job, "compensation", "")
                job_url = ""
                if job is not None:
                    job_url = getattr(job, "url", "") or getattr(job, "job_url", "") or getattr(job, "link", "")

                meta_lines = [
                    f"id: {pkg.id}",
                    f"job_title: {job_title}",
                    f"company: {company}",
                    f"location: {location}",
                    f"profile: {profile_name}",
                    f"created_at: {getattr(pkg, 'created_at', '')}",
                    f"posting_date: {posting_date}",
                    f"salary: {salary}",
                    f"job_url: {job_url}",
                ]
                zf.writestr(f"{folder}/metadata.txt", "\n".join(meta_lines))
            except Exception as exc:
                # Record the failure and continue with other packages
                pkg_id = getattr(pkg, "id", "<unknown>")
                errors.append(f"Package {pkg_id} failed: {exc}")

        # If there were errors, add a top-level failures.txt describing them
        if errors:
            zf.writestr("FAILURES.txt", "\n".join(errors))

    buf.seek(0)
    return buf.read()
