"""Resume parsing and skill extraction from PDF/DOCX files."""

import base64
import io
import json
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from .models import UserProfile

logger = logging.getLogger(__name__)


@dataclass
class ParsedResume:
    """Structured data extracted from a resume."""
    
    raw_text: str
    name: str = ""
    email: str = ""
    phone: str = ""
    summary: str = ""
    current_title: str = ""
    years_experience: Optional[int] = None
    skills: list[str] = field(default_factory=list)
    education: list[str] = field(default_factory=list)
    experience: list[dict] = field(default_factory=list)
    certifications: list[str] = field(default_factory=list)
    
    def to_user_profile(self, profile_id: Optional[str] = None) -> UserProfile:
        """Convert parsed resume to a UserProfile."""
        profile = UserProfile(
            name=self.name,
            email=self.email,
            resume_text=self.raw_text,
            summary=self.summary,
            skills=self.skills,
            years_experience=self.years_experience,
            current_title=self.current_title,
        )
        if profile_id:
            profile.id = profile_id
        return profile


class ResumeParser:
    """Parses resume files (PDF, DOCX) and extracts structured information."""
    
    def __init__(self, llm_client=None):
        """Initialize parser with optional LLM client for smart extraction.
        
        Args:
            llm_client: Azure OpenAI chat client for LLM-based extraction.
                       If None, uses regex-based extraction only.
        """
        self.llm_client = llm_client
    
    def parse_pdf(self, file_path: str) -> str:
        """Extract text from a PDF file."""
        import fitz  # PyMuPDF
        
        logger.info("[RESUME] Parsing PDF: %s", file_path)
        doc = fitz.open(file_path)
        text_parts = []
        
        for page in doc:
            text_parts.append(page.get_text())
        
        doc.close()
        return "\n".join(text_parts)
    
    def parse_pdf_bytes(self, data: bytes) -> str:
        """Extract text from PDF bytes."""
        import fitz  # PyMuPDF
        
        logger.info("[RESUME] Parsing PDF from bytes (%d bytes)", len(data))
        doc = fitz.open(stream=data, filetype="pdf")
        text_parts = []
        
        for page in doc:
            text_parts.append(page.get_text())
        
        doc.close()
        return "\n".join(text_parts)
    
    def parse_docx(self, file_path: str) -> str:
        """Extract text from a DOCX file."""
        from docx import Document
        
        logger.info("[RESUME] Parsing DOCX: %s", file_path)
        doc = Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return "\n".join(text_parts)
    
    def parse_docx_bytes(self, data: bytes) -> str:
        """Extract text from DOCX bytes."""
        from docx import Document
        
        logger.info("[RESUME] Parsing DOCX from bytes (%d bytes)", len(data))
        doc = Document(io.BytesIO(data))
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return "\n".join(text_parts)
    
    def parse_file(self, file_path: str) -> str:
        """Extract text from a file based on extension."""
        path = Path(file_path)
        extension = path.suffix.lower()
        
        if extension == ".pdf":
            return self.parse_pdf(file_path)
        elif extension in (".docx", ".doc"):
            return self.parse_docx(file_path)
        elif extension == ".txt":
            return path.read_text(encoding="utf-8")
        else:
            raise ValueError(f"Unsupported file type: {extension}")
    
    def parse_base64(self, data: str, file_type: str) -> str:
        """Parse base64-encoded file data.
        
        Args:
            data: Base64-encoded file content
            file_type: File type (pdf, docx, txt)
        """
        decoded = base64.b64decode(data)
        
        if file_type.lower() == "pdf":
            return self.parse_pdf_bytes(decoded)
        elif file_type.lower() in ("docx", "doc"):
            return self.parse_docx_bytes(decoded)
        elif file_type.lower() == "txt":
            return decoded.decode("utf-8")
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def extract_with_regex(self, text: str) -> ParsedResume:
        """Extract structured data using regex patterns.
        
        This is a fallback when LLM is not available.
        """
        parsed = ParsedResume(raw_text=text)
        
        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            parsed.email = email_match.group()
        
        # Extract phone
        phone_pattern = r'[\+]?[(]?[0-9]{1,3}[)]?[-\s\.]?[(]?[0-9]{1,4}[)]?[-\s\.]?[0-9]{1,4}[-\s\.]?[0-9]{1,9}'
        phone_match = re.search(phone_pattern, text)
        if phone_match:
            parsed.phone = phone_match.group()
        
        # Common skills to look for
        common_skills = [
            "Python", "Java", "JavaScript", "TypeScript", "C++", "C#", "Go", "Rust",
            "SQL", "PostgreSQL", "MySQL", "MongoDB", "Redis", "Elasticsearch",
            "AWS", "Azure", "GCP", "Docker", "Kubernetes", "Terraform",
            "React", "Angular", "Vue", "Node.js", "Django", "Flask", "FastAPI",
            "Machine Learning", "Deep Learning", "NLP", "Computer Vision",
            "TensorFlow", "PyTorch", "Scikit-learn", "Pandas", "NumPy",
            "Git", "CI/CD", "Jenkins", "GitHub Actions",
            "Agile", "Scrum", "Leadership", "Project Management",
            "API", "REST", "GraphQL", "Microservices",
        ]
        
        # Find skills mentioned in text
        text_lower = text.lower()
        for skill in common_skills:
            if skill.lower() in text_lower:
                parsed.skills.append(skill)
        
        # Try to extract years of experience
        exp_pattern = r'(\d+)\+?\s*(?:years?|yrs?)(?:\s+of)?\s+(?:experience|exp)'
        exp_match = re.search(exp_pattern, text, re.IGNORECASE)
        if exp_match:
            parsed.years_experience = int(exp_match.group(1))
        
        return parsed

    # --- Helper utilities for targeted matching / normalization ---
    def split_into_sections(self, text: str) -> dict:
        """Split resume text into a dict of common sections.

        Returns a mapping of lowercase section name -> section text. Heuristics are
        intentionally simple (look for common headings) to avoid over-engineering.
        """
        headings = [
            "professional summary", "summary", "experience", "work experience",
            "employment history", "skills", "key skills", "education",
            "certifications", "projects", "publications",
        ]

        hits = []
        for h in headings:
            # match heading on its own line, allow optional trailing colon
            m = re.search(rf"(?im)^\s*{re.escape(h)}\s*:?")
            if m:
                hits.append((m.start(), h))

        # If no headings found, return whole text as 'raw'
        if not hits:
            return {"raw": text.strip()}

        hits.sort()
        sections = {}
        for idx, (pos, name) in enumerate(hits):
            start = pos
            end = hits[idx + 1][0] if idx + 1 < len(hits) else len(text)
            # Extract text slice, then clean heading line if present
            slice_text = text[start:end].strip()
            # Remove the heading label if present at the top
            slice_text = re.sub(rf"(?i)^{re.escape(name)}\s*:?", "", slice_text).strip()
            sections[name.lower()] = slice_text

        return sections

    def extract_bullets(self, section_text: str) -> list:
        """Return a list of bullet lines from a section.

        Looks for common bullet markers (-, •, *) at line starts. Falls back to
        returning non-empty lines if no explicit bullets are found.
        """
        if not section_text:
            return []

        bullets = re.findall(r"(?m)^[ \t]*[-•\*]\s+(.*)$", section_text)
        if bullets:
            return [b.strip() for b in bullets if b.strip()]

        # Fallback: choose long lines (likely meaningful sentences)
        lines = [l.strip() for l in section_text.splitlines() if l.strip()]
        return [l for l in lines if len(l) > 40][:50]

    def extract_experience_entries(self, text: str) -> list:
        """Return a list of experience entry dicts extracted from the Experience section.

        Each entry is a dict with optional keys: 'company', 'title', 'dates', 'bullets', 'raw'.
        This uses simple heuristics (split on double-newline blocks) and attempts to
        parse a leading line like 'Title — Company' or 'Company — Title'.
        """
        sections = self.split_into_sections(text)
        exp_text = sections.get("experience") or sections.get("work experience") or sections.get("employment history") or text

        entries = []
        # Split entries by double-newline (common separator between roles)
        blocks = [b.strip() for b in re.split(r"\n{2,}", exp_text) if b.strip()]
        for b in blocks:
            lines = [l.strip() for l in b.splitlines() if l.strip()]
            first = lines[0] if lines else ""
            company = ""
            title = ""
            dates = ""

            # Try patterns like 'Title – Company' or 'Company — Title'
            m = re.match(r"^(?P<a>.+?)\s+[-–—]\s+(?P<b>.+)$", first)
            if m:
                # Heuristic: if first token has a comma or contains 'Inc' assume company is b
                a, bb = m.group('a').strip(), m.group('b').strip()
                if any(x in bb for x in ("Inc", "LLC", "Corp", "Company", "Ltd")) or len(bb.split()) > len(a.split()):
                    title, company = a, bb
                else:
                    company, title = a, bb
            else:
                # Try to extract dates from the first line
                d = re.search(r"(\b\d{4}\b).*[-–—to]{1,4}.*(\b\d{4}\b|Present|present)", first)
                if d:
                    dates = d.group(0)

            bullets = self.extract_bullets(b)

            entries.append({
                "company": company,
                "title": title,
                "dates": dates,
                "bullets": bullets,
                "raw": b,
            })

        return entries

    def normalize_skill_terms(self, skills: list[str]) -> list[str]:
        """Normalize skill tokens to a canonical, human-friendly form.

        This is intentionally small — expand mappings as needed.
        """
        canon = {
            "aws": "AWS",
            "amazon web services": "AWS",
            "gcp": "GCP",
            "google cloud": "GCP",
            "machine learning": "Machine Learning",
            "ml": "Machine Learning",
            "nlp": "NLP",
            "postgres": "PostgreSQL",
            "postgresql": "PostgreSQL",
            "redis": "Redis",
            "docker": "Docker",
            "kubernetes": "Kubernetes",
        }

        out = []
        for s in skills:
            key = s.strip().lower()
            mapped = canon.get(key, None)
            if mapped:
                out.append(mapped)
            else:
                # Title-case multiword tokens, preserve common all-caps
                if key.isupper() or key.lower() in ("nlp", "api", "ci/cd"):
                    out.append(s.strip())
                else:
                    out.append(s.strip().title())

        # Deduplicate while preserving order
        seen = set()
        res = []
        for s in out:
            if s not in seen:
                seen.add(s)
                res.append(s)
        return res

    def match_jd_terms(self, text: str, jd_terms: list[str]) -> list[str]:
        """Return subset of jd_terms that appear in text (case-insensitive).

        Useful to highlight which JD skills are already present in the resume.
        """
        if not text or not jd_terms:
            return []
        t = text.lower()
        found = []
        for term in jd_terms:
            if not term:
                continue
            if term.lower() in t:
                found.append(term)
        return found
    
    async def extract_with_llm(self, text: str) -> ParsedResume:
        """
        Extract structured data from any type of resume (General Purpose).
        Uses industry-agnostic logic to identify skills and experience.
        """
        if not self.llm_client:
            raise ValueError("LLM client not configured")
        
        prompt = f"""
        You are a highly accurate Resume Parser. Your goal is to extract structured data from the provided resume text, 
        regardless of the industry or profession.

        ### EXTRACTION GUIDELINES:
        1. **Skills**: Extract all relevant hard skills, soft skills, and tool proficiencies. 
           - Group them logically based on how they appear in the text.
           - DO NOT include personal contact links (e.g., LinkedIn, GitHub, Portfolio URLs) as skills.
        2. **Experience**: Identify the candidate's current or most recent job title. 
           - Calculate the total number of years of professional experience by summing the durations of all listed roles.
        3. **Summary**: Provide a concise 2-3 sentence overview of their professional background.
        4. **Education & Certifications**: Capture degrees, institutions, and any professional licenses or certificates.

        ### OUTPUT FORMAT (Strict JSON):
        {{
            "name": "Full Name",
            "email": "Email Address",
            "phone": "Phone Number",
            "summary": "Professional Summary",
            "current_title": "Most Recent Job Title",
            "years_experience": integer,
            "skills": ["A comprehensive flat list of all identified skills"],
            "education": ["List of degrees and institutions"],
            "certifications": ["List of certifications or licenses"]
        }}

        Resume text:
        ---
        {text[:8000]}
        ---
        Return ONLY valid JSON.
        """

        try:
            response = await self.llm_client.complete(
                messages=[{"role": "user", "content": prompt}],
                temperature=0.0,
                # max_tokens=2500,
            )
            
            content = response.content.strip()
            
            # Remove Markdown code blocks if the LLM includes them
            if content.startswith("```"):
                content = re.sub(r'^```json\s*|```\s*$', '', content, flags=re.MULTILINE)
            
            data = json.loads(content)
            
            return ParsedResume(
                raw_text=text,
                name=data.get("name", ""),
                email=data.get("email", ""),
                phone=data.get("phone", ""),
                summary=data.get("summary", ""),
                current_title=data.get("current_title", ""),
                years_experience=data.get("years_experience"),
                skills=data.get("skills", []),
                education=data.get("education", []),
                certifications=data.get("certifications", [])
            )
        except json.JSONDecodeError as e:
            logger.error(f"JSON parsing error: {e}")
            return self.extract_with_regex(text)
        except Exception as e:
            logger.warning(f"LLM extraction failed: {e}")
            return self.extract_with_regex(text)   
    
    
    async def parse_and_extract(
        self,
        file_path: Optional[str] = None,
        file_data: Optional[str] = None,
        file_type: Optional[str] = None,
        use_llm: bool = True,
    ) -> ParsedResume:
        """Parse a resume file and extract structured data.
        
        Args:
            file_path: Path to the resume file
            file_data: Base64-encoded file content (alternative to file_path)
            file_type: File type when using file_data (pdf, docx, txt)
            use_llm: Whether to use LLM for extraction (falls back to regex if fails)
        
        Returns:
            ParsedResume with extracted information
        """
        # Extract text from file
        if file_path:
            text = self.parse_file(file_path)
        elif file_data and file_type:
            text = self.parse_base64(file_data, file_type)
        else:
            raise ValueError("Must provide either file_path or (file_data and file_type)")
        
        # Extract structured data
        if use_llm and self.llm_client:
            return await self.extract_with_llm(text)
        else:
            return self.extract_with_regex(text)


def create_parser(llm_client=None) -> ResumeParser:
    """Factory function to create a resume parser."""
    return ResumeParser(llm_client)
